#!/usr/bin/env python3
"""
Charter School Evaluation Comments Compiler - Plain Text Version
Extracts reviewer comments from Word documents by converting to plain text.
Accepts both .txt files and .docx files (auto-converts .docx to text).
This is much more reliable than parsing Word XML structures.
"""

import os
import sys
import re
from pathlib import Path
from collections import defaultdict
from typing import Dict, List, Set, Tuple

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# Section lists for different application types
SECTION_LISTS = {
    '1': {  # Standard Application
        'name': 'Standard Application',
        'sections': [
            'Mission Guiding Principles and Purpose',
            'Target Population and Student Body',
            'Educational Program Design',
            'Curriculum and Instructional Design',
            'Student Performance',
            'Exceptional Students',
            'English Language Learners',
            'School Culture and Discipline',
            'Supplemental Programming',
            'Governance',
            'Management and Staffing',
            'Human Resources and Employment',
            'Professional Development',
            'Student Recruitment and Enrollment',
            'Parent and Community Involvement',
            'Facilities',
            'Transportation Service',
            'Food Service',
            'School Safety and Security',
            'Budget',
            'Financial Management and Oversight',
            'Start-Up Plan',
        ]
    },
    '2': {  # Virtual Application
        'name': 'Virtual Application',
        'sections': [
            'Mission, Guiding Principles and Purpose',
            'Target Population and Student Body',
            'Educational Program Design',
            'Curriculum Plan',
            'Student Performance, Assessment and Evaluation',
            'Exceptional Students',
            'English Language Learners',
            'School Culture and Discipline',
            'Supplemental Programming',
            'Governance',
            'Management and Staffing',
            'Human Resources and Employment',
            'Professional Development',
            'Student Recruitment and Enrollment',
            'Parent and Community Involvement',
            'Budget',
            'Financial Management and Oversight',
            'Start-Up Plan',
        ]
    },
    '3': {  # High Performing System Replication
        'name': 'High Performing System Replication',
        'sections': [
            'Replication Overview',
            'Mission Guiding Principles and Purpose',
            'Educational Program, Curriculum, and Instructional Design',
            'Student Performance',
            'Student Recruitment and Enrollment',
            'Management and Staffing',
            'Facilities',
            'Transportation Service',
            'Food Service',
            'School Safety and Security',
            'Budget',
            'Financial Management and Oversight',
        ]
    }
}


def extract_text_from_docx(docx_path: Path) -> str:
    """
    Extract all text from a Word document, similar to 'Save As Plain Text'.
    Preserves basic structure with tabs between table cells.
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx library is required to process .docx files. Install with: pip install python-docx")

    try:
        doc = Document(docx_path)
        text_parts = []

        # Extract all paragraphs and tables in document order
        for element in doc.element.body:
            # Check if it's a paragraph
            if element.tag.endswith('p'):
                # Find the paragraph object
                for para in doc.paragraphs:
                    if para._element == element:
                        para_text = para.text.strip()
                        if para_text:
                            text_parts.append(para_text)
                        break

            # Check if it's a table
            elif element.tag.endswith('tbl'):
                # Find the table object
                for table in doc.tables:
                    if table._element == element:
                        # Extract table content with tabs between cells
                        for row in table.rows:
                            row_cells = []
                            for cell in row.cells:
                                # Get all text from the cell
                                cell_text = cell.text.strip()
                                row_cells.append(cell_text)

                            # Join cells with tabs (similar to plain text export)
                            if any(row_cells):  # Only add non-empty rows
                                text_parts.append('\t'.join(row_cells))
                        break

        return '\n'.join(text_parts)

    except Exception as e:
        raise Exception(f"Error extracting text from {docx_path.name}: {e}")


class BoilerplateFilter:
    """Manages boilerplate text filtering."""

    def __init__(self):
        self.boilerplate_texts: Set[str] = set()
        self.boilerplate_lines: Set[str] = set()

    def load_templates(self, template_folder: Path) -> int:
        """Load all template files (.txt or .docx) and extract boilerplate text."""
        txt_files = list(template_folder.glob('*.txt'))
        docx_files = [f for f in template_folder.glob('*.docx') if not f.name.startswith('~$')]
        template_files = txt_files + docx_files

        for template_file in template_files:
            try:
                # Get content based on file type
                if template_file.suffix.lower() == '.docx':
                    if not DOCX_AVAILABLE:
                        print(f"  Warning: Skipping {template_file.name} (python-docx not installed)")
                        continue
                    content = extract_text_from_docx(template_file)
                else:
                    content = template_file.read_text(encoding='utf-8', errors='ignore')

                # Extract all lines
                for line in content.split('\n'):
                    clean_line = line.strip()
                    if clean_line:
                        self.boilerplate_texts.add(clean_line)
                        self.boilerplate_lines.add(clean_line)

            except Exception as e:
                print(f"  Warning: Could not load template {template_file.name}: {e}")

        return len(template_files)

    def is_boilerplate(self, text: str) -> bool:
        """Check if text is boilerplate."""
        text = text.strip()
        if not text:
            return True

        # Check exact match
        if text in self.boilerplate_texts or text in self.boilerplate_lines:
            return True

        # Check if it's a common header or placeholder
        lower_text = text.lower()
        if lower_text in ['strengths', 'concerns and additional questions',
                          'concerns', 'reference', 'references',
                          'meets the standard', 'partially meets the standard',
                          'does not meet the standard', 'type', 'comment',
                          'strength', 'concern', 'question', 'follow up',
                          'choose type', 'enter comment here.', 'enter comment here']:
            return True

        return False


class CommentExtractor:
    """Extracts comments from plain text evaluation documents."""

    def __init__(self, boilerplate_filter: BoilerplateFilter, section_names: List[str]):
        self.boilerplate_filter = boilerplate_filter
        self.section_names = section_names
        self.comments: Dict[str, Dict[str, List[Tuple[str, str]]]] = defaultdict(
            lambda: {'strengths': [], 'concerns': []}
        )
        self.boilerplate_filtered_count = 0

    def extract_reviewer_name(self, content: str) -> str:
        """Extract reviewer name from document content."""
        # Look for "Reviewer's Name" followed by the name
        match = re.search(r"Reviewer[''']?s?\s+Name\s*\n\s*(.+)", content, re.IGNORECASE)
        if match:
            name = match.group(1).strip()
            # Sometimes the date appears on the same line or next line
            name = re.split(r'\d{1,2}/\d{1,2}/\d{2,4}', name)[0].strip()
            if name and name not in ['Review Team Initial Date', 'Mark Cannon']:
                return name

        # Fallback: try to find it in a table format
        lines = content.split('\n')
        for i, line in enumerate(lines):
            if "reviewer" in line.lower() and "name" in line.lower():
                # Check next few lines for the actual name
                for j in range(i+1, min(i+5, len(lines))):
                    potential_name = lines[j].strip()
                    if potential_name and not re.match(r'\d', potential_name) and len(potential_name) > 2:
                        return potential_name

        return 'Unknown Reviewer'

    def extract_school_name(self, content: str) -> str:
        """Extract school name from document content."""
        # Look for "Proposed Charter School Name"
        match = re.search(r"Proposed\s+Charter\s+School\s+Name\s*\n\s*(.+)", content, re.IGNORECASE)
        if match:
            return match.group(1).strip()

        # Look for other patterns
        match = re.search(r"School\s+Name[:\s]+(.+)", content, re.IGNORECASE)
        if match:
            return match.group(1).strip()

        return 'Unknown School'

    def find_section_name(self, text: str) -> str:
        """Find matching section name from text."""
        text = text.strip()

        # Try to match against known section names
        for section_name in self.section_names:
            # Check if section name appears in text (case-insensitive)
            if section_name.lower() in text.lower():
                return section_name

            # Check if text appears in section name (partial match)
            if len(text) > 10 and text.lower() in section_name.lower():
                return section_name

        # Check for addendum sections
        if 'addendum' in text.lower():
            return text

        return None

    def parse_section_comments(self, section_text: str) -> Tuple[List[Tuple[str, str]], List[Tuple[str, str]]]:
        """Parse a section to extract strengths and concerns.
        Returns: (strengths, concerns) where each is a list of (reference, comment) tuples.
        """
        strengths = []
        concerns = []

        # Split into lines
        lines = section_text.split('\n')

        # Look for table data in Type | Reference | Comment format
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            # Try to detect table row with type, reference, comment
            # Patterns: "Strength\t10\tMission is..." or "Strength  10  Mission is..."
            parts = re.split(r'\t+|\s{2,}', line, maxsplit=2)

            if len(parts) >= 3:
                comment_type = parts[0].strip().lower()
                reference = parts[1].strip()
                comment = parts[2].strip()

                # Skip if this looks like a header row
                if 'type' in comment_type or comment_type in ['strength', 'concern', 'question', 'follow up']:
                    if 'reference' in reference.lower() or 'comment' in comment.lower():
                        continue

                # Skip boilerplate
                if self.boilerplate_filter.is_boilerplate(comment):
                    if comment:
                        self.boilerplate_filtered_count += 1
                    continue

                # Add page reference if available and not boilerplate
                comment_text = comment
                if reference and not self.boilerplate_filter.is_boilerplate(reference):
                    # Check if reference looks like a page number
                    if reference.replace('.', '').replace('p', '').replace(' ', '').isdigit():
                        page_num = reference.replace('p.', '').replace('p', '').strip()
                        comment_text = f"{comment} [p. {page_num}]"
                    elif len(reference) < 20:  # Short reference
                        comment_text = f"{comment} [{reference}]"

                # Categorize by type
                if 'strength' in comment_type:
                    strengths.append((reference, comment_text))
                elif 'concern' in comment_type or 'question' in comment_type:
                    concerns.append((reference, comment_text))

        return strengths, concerns

    def process_document(self, file_path: Path) -> Tuple[int, str, str]:
        """Process a single evaluation document (.txt or .docx).
        Returns: (comment_count, reviewer_name, school_name)
        """
        try:
            # Read or extract the text content
            if file_path.suffix.lower() == '.docx':
                # Extract text from Word document
                content = extract_text_from_docx(file_path)
            else:
                # Read plain text file
                content = file_path.read_text(encoding='utf-8', errors='ignore')

            # Extract metadata
            reviewer_name = self.extract_reviewer_name(content)
            school_name = self.extract_school_name(content)

            comment_count = 0

            # Find section evaluation comments blocks
            # Pattern: "Section X Evaluation Comments:" or "Section XX (YY). Title"
            section_pattern = r'(?:Section\s+\d+.*?Evaluation\s+Comments:|Addendum\s+[A-Z]\s+Evaluation\s+Comments:)'

            # Split content into sections
            section_blocks = re.split(section_pattern, content)
            section_headers = re.findall(section_pattern, content)

            # Match sections to known section names
            for idx, section_block in enumerate(section_blocks[1:], 0):  # Skip first block (header content)
                if idx >= len(section_headers):
                    break

                header = section_headers[idx]

                # Try to identify which section this is
                current_section = None

                # Look backwards in content to find the section title
                header_pos = content.find(header)
                if header_pos > 0:
                    # Get previous 500 characters to find section title
                    context = content[max(0, header_pos - 500):header_pos]
                    context_lines = context.split('\n')

                    # Look for section numbers and titles
                    for line in reversed(context_lines[-10:]):
                        line = line.strip()
                        # Try patterns like "1. Mission" or "10. Governance"
                        match = re.match(r'^\s*\d+\.?\s+(.+)', line)
                        if match:
                            potential_title = match.group(1).strip()
                            current_section = self.find_section_name(potential_title)
                            if current_section:
                                break
                        else:
                            # Try direct match
                            current_section = self.find_section_name(line)
                            if current_section:
                                break

                # Check if this is an Addendum section
                if not current_section and 'addendum' in header.lower():
                    # Extract addendum title from context
                    header_pos = content.find(header)
                    if header_pos > 0:
                        context = content[max(0, header_pos - 300):header_pos]
                        context_lines = [l.strip() for l in context.split('\n') if l.strip()]
                        if context_lines:
                            current_section = context_lines[-1]

                if not current_section:
                    current_section = "General Comments"

                # Parse comments from this section
                strengths, concerns = self.parse_section_comments(section_block)

                # Add to comments dictionary
                for ref, strength in strengths:
                    if not self.boilerplate_filter.is_boilerplate(strength):
                        self.comments[current_section]['strengths'].append(
                            (reviewer_name, strength)
                        )
                        comment_count += 1

                for ref, concern in concerns:
                    if not self.boilerplate_filter.is_boilerplate(concern):
                        self.comments[current_section]['concerns'].append(
                            (reviewer_name, concern)
                        )
                        comment_count += 1

            return comment_count, reviewer_name, school_name

        except Exception as e:
            print(f"  ✗ Error processing {file_path.name}: {e}")
            import traceback
            traceback.print_exc()
            return 0, 'Unknown Reviewer', 'Unknown School'

    def remove_duplicates(self) -> int:
        """Remove exact duplicate comments."""
        duplicates_removed = 0

        for section in self.comments:
            # Remove duplicates from strengths
            seen = set()
            unique_strengths = []
            for reviewer, comment in self.comments[section]['strengths']:
                if comment not in seen:
                    seen.add(comment)
                    unique_strengths.append((reviewer, comment))
                else:
                    duplicates_removed += 1
            self.comments[section]['strengths'] = unique_strengths

            # Remove duplicates from concerns
            seen = set()
            unique_concerns = []
            for reviewer, comment in self.comments[section]['concerns']:
                if comment not in seen:
                    seen.add(comment)
                    unique_concerns.append((reviewer, comment))
                else:
                    duplicates_removed += 1
            self.comments[section]['concerns'] = unique_concerns

        return duplicates_removed

    def generate_markdown(self, school_name: str) -> str:
        """Generate Markdown output."""
        lines = []
        lines.append("# Charter Application Review Comments Compilation")
        lines.append(f"## School Name: {school_name}")
        lines.append("")

        # Process sections in order
        for idx, section_name in enumerate(self.section_names, 1):
            if section_name not in self.comments:
                continue

            section_data = self.comments[section_name]

            # Skip if no comments in this section
            if not section_data['strengths'] and not section_data['concerns']:
                continue

            lines.append(f"## Section {idx}. {section_name}")

            # Add strengths
            if section_data['strengths']:
                lines.append("### Strengths")
                for reviewer, comment in section_data['strengths']:
                    lines.append(f"- {reviewer}: {comment}")
                lines.append("")

            # Add concerns
            if section_data['concerns']:
                lines.append("### Concerns")
                for reviewer, comment in section_data['concerns']:
                    lines.append(f"- {reviewer}: {comment}")
                lines.append("")

        # Add any addendum sections or other sections not in the main list
        for section_name in self.comments:
            if section_name not in self.section_names:
                section_data = self.comments[section_name]

                if not section_data['strengths'] and not section_data['concerns']:
                    continue

                lines.append(f"## {section_name}")

                if section_data['strengths']:
                    lines.append("### Strengths")
                    for reviewer, comment in section_data['strengths']:
                        lines.append(f"- {reviewer}: {comment}")
                    lines.append("")

                if section_data['concerns']:
                    lines.append("### Concerns")
                    for reviewer, comment in section_data['concerns']:
                        lines.append(f"- {reviewer}: {comment}")
                    lines.append("")

        return '\n'.join(lines)


def main():
    """Main program execution."""
    print("Charter School Evaluation Comments Compiler - Plain Text Version")
    print("=" * 66)
    print()
    print("This version accepts both .docx and .txt files.")
    print("Word files (.docx) are automatically converted to text.")
    print("You can also manually export: File → Save As → Plain Text (.txt)")
    print()

    # Prompt for review documents folder
    review_folder_path = input("Enter path to folder containing review files (.txt or .docx): ").strip()
    review_folder = Path(review_folder_path)

    if not review_folder.exists() or not review_folder.is_dir():
        print(f"Error: Folder not found: {review_folder_path}")
        sys.exit(1)

    # Find both .txt and .docx files
    txt_files = [f for f in review_folder.glob('*.txt') if not f.name.startswith('~$')]
    docx_files = [f for f in review_folder.glob('*.docx') if not f.name.startswith('~$')]
    all_files = txt_files + docx_files

    if not all_files:
        print(f"Error: No .txt or .docx files found in {review_folder_path}")
        sys.exit(1)

    # Show what we found
    if txt_files and docx_files:
        print(f"Found {len(txt_files)} .txt file(s) and {len(docx_files)} .docx file(s)")
    elif txt_files:
        print(f"Found {len(txt_files)} .txt file(s)")
    else:
        print(f"Found {len(docx_files)} .docx file(s)")

    if docx_files and not DOCX_AVAILABLE:
        print()
        print("WARNING: python-docx library not installed.")
        print("Install with: pip install python-docx")
        print("Only .txt files will be processed.")
        all_files = txt_files
        if not all_files:
            print("Error: No processable files found")
            sys.exit(1)

    print()

    # Prompt for template folder (optional)
    template_folder_path = input("Enter path to templates folder (or press Enter to skip): ").strip()

    boilerplate_filter = BoilerplateFilter()

    if template_folder_path:
        template_folder = Path(template_folder_path)
        if template_folder.exists() and template_folder.is_dir():
            print("Loading templates...")
            template_count = boilerplate_filter.load_templates(template_folder)
            print(f"Loaded {template_count} template file(s)")
        else:
            print("Template folder not found, proceeding without template filtering")

    print()

    # Prompt for application type
    print("Select application type:")
    print("  1. Standard Application (22 sections)")
    print("  2. Virtual Application (18 sections)")
    print("  3. High Performing System Replication (12 sections)")

    app_type = input("Enter selection (1-3): ").strip()

    if app_type not in SECTION_LISTS:
        print(f"Error: Invalid selection: {app_type}")
        sys.exit(1)

    section_info = SECTION_LISTS[app_type]
    print()

    # Process documents
    print("Processing reviews...")
    extractor = CommentExtractor(boilerplate_filter, section_info['sections'])

    school_name = None

    for review_file in all_files:
        comment_count, reviewer_name, file_school_name = extractor.process_document(review_file)
        if not school_name:
            school_name = file_school_name
        print(f"  ✓ {review_file.name} - {reviewer_name} ({comment_count} comments extracted)")

    if not school_name:
        school_name = "Unknown School"

    print()

    # Remove duplicates
    duplicates = extractor.remove_duplicates()

    # Calculate total unique comments
    total_comments = sum(
        len(section['strengths']) + len(section['concerns'])
        for section in extractor.comments.values()
    )

    print(f"Removed {extractor.boilerplate_filtered_count} boilerplate text instances")
    print(f"Removed {duplicates} duplicate comments")
    print(f"Compiled {total_comments} unique comments across {len(extractor.comments)} sections")
    print()

    # Prompt for output filename
    default_output = f"{school_name}_CompiledReviews.md"
    output_filename = input(f"Enter output filename [{default_output}]: ").strip()
    if not output_filename:
        output_filename = default_output

    output_path = review_folder / output_filename

    # Generate and save output
    markdown_output = extractor.generate_markdown(school_name)
    output_path.write_text(markdown_output, encoding='utf-8')

    print(f"✓ Output saved to: {output_path}")


if __name__ == '__main__':
    main()
