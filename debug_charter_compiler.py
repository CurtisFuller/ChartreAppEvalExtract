#!/usr/bin/env python3
"""
Debug tool for Charter School Evaluation Comments Compiler
Analyzes Word documents and outputs detailed structure information for troubleshooting.
"""

import os
import sys
from pathlib import Path
from datetime import datetime
from docx import Document


def extract_text_from_cell(cell):
    """Extract text from a cell, including form fields."""
    # Try regular text first
    text = cell.text.strip()

    # Check for form fields in the cell
    try:
        from docx.oxml.ns import qn
        # Look for form field elements
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                # Check for form field in run element
                elem = run._element
                for child in elem.iter():
                    tag = child.tag
                    if 'ffData' in tag or 'textInput' in tag:
                        # This indicates a form field
                        return text, True
    except:
        pass

    return text, False


def analyze_document(file_path: Path, debug_file, boilerplate_texts=None):
    """Analyze a single Word document and write debug info."""
    debug_file.write(f"\n{'=' * 80}\n")
    debug_file.write(f"FILE: {file_path.name}\n")
    debug_file.write(f"{'=' * 80}\n\n")

    try:
        doc = Document(file_path)

        # Analyze paragraphs
        debug_file.write(f"PARAGRAPHS ({len(doc.paragraphs)} total):\n")
        debug_file.write("-" * 80 + "\n")

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:  # Only show non-empty paragraphs
                debug_file.write(f"\n[Para {idx}] Style: {para.style.name}\n")
                debug_file.write(f"Text: {text}\n")

                # Check if it might be a section header
                if any(keyword in text.lower() for keyword in ['section', 'mission', 'target', 'educational', 'curriculum']):
                    debug_file.write(f"*** POTENTIAL SECTION HEADER ***\n")

        # Analyze tables
        debug_file.write(f"\n\nTABLES ({len(doc.tables)} total):\n")
        debug_file.write("-" * 80 + "\n")

        for table_idx, table in enumerate(doc.tables):
            debug_file.write(f"\n[Table {table_idx}] Dimensions: {len(table.rows)} rows x {len(table.columns)} columns\n")
            debug_file.write("-" * 40 + "\n")

            for row_idx, row in enumerate(table.rows):
                debug_file.write(f"\n  [Row {row_idx}] {len(row.cells)} cells:\n")

                for cell_idx, cell in enumerate(row.cells):
                    text, has_form_field = extract_text_from_cell(cell)

                    # Show cell content
                    debug_file.write(f"    [Cell {cell_idx}]: ")

                    if has_form_field:
                        debug_file.write("[FORM FIELD] ")

                    if not text:
                        debug_file.write("(EMPTY)\n")
                    else:
                        # Truncate very long text for readability
                        if len(text) > 200:
                            display_text = text[:200] + "... [TRUNCATED]"
                        else:
                            display_text = text

                        debug_file.write(f"{repr(display_text)}\n")

                        # Check if it's a known header
                        text_lower = text.lower()
                        if 'strength' in text_lower:
                            debug_file.write(f"      *** IDENTIFIED AS: STRENGTHS HEADER ***\n")
                        elif 'concern' in text_lower or 'question' in text_lower:
                            debug_file.write(f"      *** IDENTIFIED AS: CONCERNS HEADER ***\n")
                        elif text_lower in ['reference', 'references']:
                            debug_file.write(f"      *** IDENTIFIED AS: REFERENCE HEADER ***\n")
                        elif 'meet' in text_lower and 'standard' in text_lower:
                            debug_file.write(f"      *** IDENTIFIED AS: STANDARDS RATING ***\n")

                        # Check if it's boilerplate
                        if boilerplate_texts and text in boilerplate_texts:
                            debug_file.write(f"      *** MARKED AS: BOILERPLATE ***\n")
                        elif text_lower in ['strengths', 'concerns and additional questions',
                                           'concerns', 'reference', 'references']:
                            debug_file.write(f"      *** MARKED AS: STANDARD HEADER ***\n")
                        elif len(text) > 50:  # Only mark potential comments if long enough
                            debug_file.write(f"      *** POTENTIAL COMMENT (length: {len(text)}) ***\n")

            debug_file.write("\n")

    except Exception as e:
        debug_file.write(f"\nERROR analyzing document: {e}\n")
        import traceback
        debug_file.write(traceback.format_exc())


def analyze_templates(template_folder: Path, debug_file):
    """Analyze template files and show boilerplate text."""
    debug_file.write(f"\n{'=' * 80}\n")
    debug_file.write(f"TEMPLATE ANALYSIS\n")
    debug_file.write(f"{'=' * 80}\n\n")

    template_files = list(template_folder.glob('*.docx'))
    debug_file.write(f"Template files found: {len(template_files)}\n\n")

    boilerplate_texts = set()
    boilerplate_lines = set()

    for template_file in template_files:
        debug_file.write(f"\nTemplate: {template_file.name}\n")
        debug_file.write("-" * 40 + "\n")

        try:
            doc = Document(template_file)

            # Extract all paragraph text
            para_count = 0
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    boilerplate_texts.add(text)
                    para_count += 1
                    for line in text.split('\n'):
                        clean_line = line.strip()
                        if clean_line:
                            boilerplate_lines.add(clean_line)

            # Extract all table text
            table_count = 0
            cell_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            boilerplate_texts.add(text)
                            cell_count += 1
                            for line in text.split('\n'):
                                clean_line = line.strip()
                                if clean_line:
                                    boilerplate_lines.add(clean_line)

            debug_file.write(f"  Paragraphs with text: {para_count}\n")
            debug_file.write(f"  Tables: {table_count}\n")
            debug_file.write(f"  Table cells with text: {cell_count}\n")

        except Exception as e:
            debug_file.write(f"  ERROR: {e}\n")

    debug_file.write(f"\nTotal boilerplate texts collected: {len(boilerplate_texts)}\n")
    debug_file.write(f"Total boilerplate lines collected: {len(boilerplate_lines)}\n")

    # Show sample boilerplate (first 20 items)
    debug_file.write(f"\nSample boilerplate texts (first 20):\n")
    debug_file.write("-" * 40 + "\n")
    for idx, text in enumerate(list(boilerplate_texts)[:20]):
        if len(text) > 100:
            display_text = text[:100] + "..."
        else:
            display_text = text
        debug_file.write(f"{idx + 1}. {repr(display_text)}\n")

    return boilerplate_texts, boilerplate_lines


def main():
    """Main debug program execution."""
    print("Charter School Evaluation Comments Compiler - DEBUG MODE")
    print("=" * 57)
    print()

    # Prompt for review documents folder
    review_folder_path = input("Enter path to folder containing review documents: ").strip()
    review_folder = Path(review_folder_path)

    if not review_folder.exists() or not review_folder.is_dir():
        print(f"Error: Folder not found: {review_folder_path}")
        sys.exit(1)

    docx_files = list(review_folder.glob('*.docx'))
    # Filter out temp files
    docx_files = [f for f in docx_files if not f.name.startswith('~$')]

    if not docx_files:
        print(f"Error: No .docx files found in {review_folder_path}")
        sys.exit(1)

    print(f"Found {len(docx_files)} .docx file(s)")
    print()

    # Prompt for template folder
    template_folder_path = input("Enter path to folder containing evaluation matrix templates: ").strip()
    template_folder = Path(template_folder_path)

    if not template_folder.exists() or not template_folder.is_dir():
        print(f"Error: Folder not found: {template_folder_path}")
        sys.exit(1)

    # Create debug output file
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    debug_filename = f"debug_output_{timestamp}.txt"
    debug_path = review_folder / debug_filename

    print(f"Creating debug output file: {debug_filename}")
    print()

    with open(debug_path, 'w', encoding='utf-8') as debug_file:
        # Write header
        debug_file.write("CHARTER REVIEW COMPILER - DEBUG OUTPUT\n")
        debug_file.write(f"Generated: {datetime.now()}\n")
        debug_file.write(f"Review folder: {review_folder}\n")
        debug_file.write(f"Template folder: {template_folder}\n")
        debug_file.write(f"Files to analyze: {len(docx_files)}\n")
        debug_file.write("\n")

        # Analyze templates first
        print("Analyzing templates...")
        boilerplate_texts, boilerplate_lines = analyze_templates(template_folder, debug_file)

        # Analyze each review document
        print("Analyzing review documents...")
        for idx, docx_file in enumerate(docx_files, 1):
            print(f"  [{idx}/{len(docx_files)}] {docx_file.name}")
            analyze_document(docx_file, debug_file, boilerplate_texts)

        # Write summary
        debug_file.write(f"\n{'=' * 80}\n")
        debug_file.write(f"ANALYSIS COMPLETE\n")
        debug_file.write(f"{'=' * 80}\n")
        debug_file.write(f"Total files analyzed: {len(docx_files)}\n")
        debug_file.write(f"Template boilerplate items: {len(boilerplate_texts)}\n")

    print()
    print(f"Debug output saved to: {debug_path}")
    print()
    print("Review the debug file to see:")
    print("  - Document structure (paragraphs and tables)")
    print("  - Table contents and cell values")
    print("  - Potential section headers")
    print("  - Potential comment text")
    print("  - Boilerplate text identification")
    print()
    print("Look for:")
    print("  1. Are section headers being detected?")
    print("  2. Are 'Strengths' and 'Concerns' headers found in tables?")
    print("  3. Is comment text marked as 'POTENTIAL COMMENT'?")
    print("  4. Is actual comment text being marked as BOILERPLATE incorrectly?")


if __name__ == '__main__':
    main()
