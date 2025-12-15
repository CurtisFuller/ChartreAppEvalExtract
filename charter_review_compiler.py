#!/usr/bin/env python3
"""
Charter School Evaluation Comments Compiler
Extracts reviewer comments from Word documents and compiles them into organized Markdown.
"""

import os
import sys
from pathlib import Path
from collections import defaultdict
from typing import Dict, List, Set, Tuple
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


# Section lists for different application types
SECTION_LISTS = {
    '1': {  # Standard Application
        'name': 'Standard Application',
        'sections': [
            'Mission Guiding Principles and Purpose',
            'Target Population and Student Body',
            'Educational Program Design',
            'Curriculum and Instructional Design',
            'Student Performance',
            'Exceptional Students',
            'English Language Learners',
            'School Culture and Discipline',
            'Supplemental Programming',
            'Governance',
            'Management and Staffing',
            'Human Resources and Employment',
            'Professional Development',
            'Student Recruitment and Enrollment',
            'Parent and Community Involvement',
            'Facilities',
            'Transportation Service',
            'Food Service',
            'School Safety and Security',
            'Budget',
            'Financial Management and Oversight',
            'Start-Up Plan',
        ]
    },
    '2': {  # Virtual Application
        'name': 'Virtual Application',
        'sections': [
            'Mission Guiding Principles and Purpose',
            'Target Population and Student Body',
            'Educational Program Design',
            'Curriculum and Instructional Design',
            'Student Performance',
            'Exceptional Students',
            'English Language Learners',
            'School Culture and Discipline',
            'Supplemental Programming',
            'Governance',
            'Management and Staffing',
            'Human Resources and Employment',
            'Professional Development',
            'Student Recruitment and Enrollment',
            'Parent and Community Involvement',
            'Budget',
            'Financial Management and Oversight',
            'Start-Up Plan',
        ]
    },
    '3': {  # High Performing System Replication
        'name': 'High Performing System Replication',
        'sections': [
            'Replication Overview',
            'Mission Guiding Principles and Purpose',
            'Educational Program, Curriculum, and Instructional Design',
            'Student Performance',
            'Student Recruitment and Enrollment',
            'Management and Staffing',
            'Facilities',
            'Transportation Service',
            'Food Service',
            'School Safety and Security',
            'Budget',
            'Financial Management and Oversight',
        ]
    }
}


class BoilerplateFilter:
    """Manages boilerplate text extraction from templates."""

    def __init__(self):
        self.boilerplate_texts: Set[str] = set()
        self.boilerplate_lines: Set[str] = set()

    def load_templates(self, template_folder: Path) -> int:
        """Load all template files and extract boilerplate text."""
        template_files = list(template_folder.glob('*.docx'))

        for template_file in template_files:
            try:
                doc = Document(template_file)

                # Extract all paragraph text
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        self.boilerplate_texts.add(text)
                        # Also store individual lines
                        for line in text.split('\n'):
                            clean_line = line.strip()
                            if clean_line:
                                self.boilerplate_lines.add(clean_line)

                # Extract all table text
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text = cell.text.strip()
                            if text:
                                self.boilerplate_texts.add(text)
                                for line in text.split('\n'):
                                    clean_line = line.strip()
                                    if clean_line:
                                        self.boilerplate_lines.add(clean_line)
            except Exception as e:
                print(f"  Warning: Could not load template {template_file.name}: {e}")

        return len(template_files)

    def is_boilerplate(self, text: str) -> bool:
        """Check if text is boilerplate."""
        text = text.strip()
        if not text:
            return True

        # Check exact match
        if text in self.boilerplate_texts or text in self.boilerplate_lines:
            return True

        # Check if it's a common header
        lower_text = text.lower()
        if lower_text in ['strengths', 'concerns and additional questions',
                          'concerns', 'reference', 'references',
                          'meets the standard', 'partially meets the standard',
                          'does not meet the standard']:
            return True

        return False


class CommentExtractor:
    """Extracts comments from evaluation documents."""

    def __init__(self, boilerplate_filter: BoilerplateFilter, section_names: List[str]):
        self.boilerplate_filter = boilerplate_filter
        self.section_names = section_names
        self.comments: Dict[str, Dict[str, List[Tuple[str, str]]]] = defaultdict(
            lambda: {'strengths': [], 'concerns': []}
        )
        self.boilerplate_filtered_count = 0

    def extract_reviewer_name(self, filename: str) -> str:
        """Extract reviewer name from filename."""
        # Format: SchoolName_Eval_ReviewerName.docx
        parts = filename.replace('.docx', '').split('_')
        if len(parts) >= 3 and parts[-2].lower() == 'eval':
            return parts[-1]
        return 'Unknown Reviewer'

    def extract_school_name(self, filename: str) -> str:
        """Extract school name from filename."""
        parts = filename.replace('.docx', '').split('_')
        if len(parts) >= 3 and parts[-2].lower() == 'eval':
            return '_'.join(parts[:-2])
        return 'Unknown School'

    def find_section_name(self, text: str) -> str:
        """Find matching section name from text."""
        text = text.strip()

        # Try to match against known section names
        for section_name in self.section_names:
            # Check if section name appears in text (case-insensitive)
            if section_name.lower() in text.lower():
                return section_name

            # Check if text appears in section name
            if text.lower() in section_name.lower():
                return section_name

        # Check for addendum sections
        if 'addendum' in text.lower():
            return text

        return None

    def parse_table_for_comments(self, table: Table) -> Tuple[List[str], List[str]]:
        """Parse a table to extract strengths and concerns."""
        strengths = []
        concerns = []

        current_section = None

        for row_idx, row in enumerate(table.rows):
            # Get cell texts
            cells_text = [cell.text.strip() for cell in row.cells]

            # Check if this is a header row
            first_cell = cells_text[0].lower() if cells_text else ''

            if 'strength' in first_cell:
                current_section = 'strengths'
                continue
            elif 'concern' in first_cell or 'question' in first_cell:
                current_section = 'concerns'
                continue

            # Extract content from cells
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()

                # Skip if boilerplate or empty
                if not text or self.boilerplate_filter.is_boilerplate(text):
                    continue

                # Skip if it looks like a header
                if text.lower() in ['reference', 'references', 'strengths',
                                   'concerns', 'concerns and additional questions']:
                    continue

                # Extract page references if in last column
                comment_text = text
                if cell_idx == len(row.cells) - 1 and cell_idx > 0:
                    # This might be a reference column, skip it
                    # But check if it looks like a page number
                    if text.startswith('p.') or text.isdigit():
                        continue

                # Check if comment contains page reference at the end
                # Format could be [p. 23] or (p. 23) or just p. 23

                # Add to appropriate section
                if current_section == 'strengths':
                    strengths.append(comment_text)
                elif current_section == 'concerns':
                    concerns.append(comment_text)

        return strengths, concerns

    def process_document(self, file_path: Path) -> int:
        """Process a single evaluation document."""
        try:
            doc = Document(file_path)
            reviewer_name = self.extract_reviewer_name(file_path.name)
            comment_count = 0

            current_section = None

            # First, try to identify sections from paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Check if this is a section header
                section_name = self.find_section_name(text)
                if section_name:
                    current_section = section_name

            # Process tables
            for table in doc.tables:
                strengths, concerns = self.parse_table_for_comments(table)

                # If we have comments but no current section, try to find it
                if (strengths or concerns) and not current_section:
                    # Look for section context in surrounding paragraphs
                    # For now, use a generic section
                    current_section = "General Comments"

                if current_section:
                    for strength in strengths:
                        if not self.boilerplate_filter.is_boilerplate(strength):
                            self.comments[current_section]['strengths'].append(
                                (reviewer_name, strength)
                            )
                            comment_count += 1
                        else:
                            self.boilerplate_filtered_count += 1

                    for concern in concerns:
                        if not self.boilerplate_filter.is_boilerplate(concern):
                            self.comments[current_section]['concerns'].append(
                                (reviewer_name, concern)
                            )
                            comment_count += 1
                        else:
                            self.boilerplate_filtered_count += 1

            return comment_count

        except Exception as e:
            print(f"  ✗ Error processing {file_path.name}: {e}")
            return 0

    def remove_duplicates(self) -> int:
        """Remove exact duplicate comments."""
        duplicates_removed = 0

        for section in self.comments:
            # Remove duplicates from strengths
            seen = set()
            unique_strengths = []
            for reviewer, comment in self.comments[section]['strengths']:
                if comment not in seen:
                    seen.add(comment)
                    unique_strengths.append((reviewer, comment))
                else:
                    duplicates_removed += 1
            self.comments[section]['strengths'] = unique_strengths

            # Remove duplicates from concerns
            seen = set()
            unique_concerns = []
            for reviewer, comment in self.comments[section]['concerns']:
                if comment not in seen:
                    seen.add(comment)
                    unique_concerns.append((reviewer, comment))
                else:
                    duplicates_removed += 1
            self.comments[section]['concerns'] = unique_concerns

        return duplicates_removed

    def generate_markdown(self, school_name: str) -> str:
        """Generate Markdown output."""
        lines = []
        lines.append("# Charter Application Review Comments Compilation")
        lines.append(f"## School Name: {school_name}")
        lines.append("")

        # Process sections in order
        for idx, section_name in enumerate(self.section_names, 1):
            if section_name not in self.comments:
                continue

            section_data = self.comments[section_name]

            # Skip if no comments in this section
            if not section_data['strengths'] and not section_data['concerns']:
                continue

            lines.append(f"## Section {idx}. {section_name}")

            # Add strengths
            if section_data['strengths']:
                lines.append("### Strengths")
                for reviewer, comment in section_data['strengths']:
                    lines.append(f"- {reviewer}: {comment}")
                lines.append("")

            # Add concerns
            if section_data['concerns']:
                lines.append("### Concerns")
                for reviewer, comment in section_data['concerns']:
                    lines.append(f"- {reviewer}: {comment}")
                lines.append("")

        # Add any addendum sections
        for section_name in self.comments:
            if section_name not in self.section_names and section_name != "General Comments":
                section_data = self.comments[section_name]

                if not section_data['strengths'] and not section_data['concerns']:
                    continue

                lines.append(f"## {section_name}")

                if section_data['strengths']:
                    lines.append("### Strengths")
                    for reviewer, comment in section_data['strengths']:
                        lines.append(f"- {reviewer}: {comment}")
                    lines.append("")

                if section_data['concerns']:
                    lines.append("### Concerns")
                    for reviewer, comment in section_data['concerns']:
                        lines.append(f"- {reviewer}: {comment}")
                    lines.append("")

        return '\n'.join(lines)


def main():
    """Main program execution."""
    print("Charter School Evaluation Comments Compiler")
    print("=" * 43)
    print()

    # Prompt for review documents folder
    review_folder_path = input("Enter path to folder containing review documents: ").strip()
    review_folder = Path(review_folder_path)

    if not review_folder.exists() or not review_folder.is_dir():
        print(f"Error: Folder not found: {review_folder_path}")
        sys.exit(1)

    docx_files = list(review_folder.glob('*.docx'))
    if not docx_files:
        print(f"Error: No .docx files found in {review_folder_path}")
        sys.exit(1)

    print(f"Found {len(docx_files)} .docx file(s)")
    print()

    # Prompt for template folder
    template_folder_path = input("Enter path to folder containing evaluation matrix templates: ").strip()
    template_folder = Path(template_folder_path)

    if not template_folder.exists() or not template_folder.is_dir():
        print(f"Error: Folder not found: {template_folder_path}")
        sys.exit(1)

    # Load templates
    print("Loading templates...")
    boilerplate_filter = BoilerplateFilter()
    template_count = boilerplate_filter.load_templates(template_folder)
    print(f"Loaded {template_count} template file(s)")
    print()

    # Prompt for application type
    print("Select application type:")
    print("  1. Standard Application (22 sections)")
    print("  2. Virtual Application (18 sections)")
    print("  3. High Performing System Replication (12 sections)")

    app_type = input("Enter selection (1-3): ").strip()

    if app_type not in SECTION_LISTS:
        print(f"Error: Invalid selection: {app_type}")
        sys.exit(1)

    section_info = SECTION_LISTS[app_type]
    print()

    # Extract school name from first file
    school_name = CommentExtractor(boilerplate_filter, section_info['sections']).extract_school_name(docx_files[0].name)

    # Prompt for output filename
    default_output = f"{school_name}_CompiledReviews.md"
    output_filename = input(f"Enter output filename [{default_output}]: ").strip()
    if not output_filename:
        output_filename = default_output

    output_path = review_folder / output_filename
    print()

    # Process documents
    print("Processing reviews...")
    extractor = CommentExtractor(boilerplate_filter, section_info['sections'])

    for docx_file in docx_files:
        comment_count = extractor.process_document(docx_file)
        print(f"  ✓ {docx_file.name} ({comment_count} comments extracted)")

    print()

    # Remove duplicates
    duplicates = extractor.remove_duplicates()

    # Calculate total unique comments
    total_comments = sum(
        len(section['strengths']) + len(section['concerns'])
        for section in extractor.comments.values()
    )

    print(f"Removed {extractor.boilerplate_filtered_count} boilerplate text instances")
    print(f"Removed {duplicates} duplicate comments")
    print(f"Compiled {total_comments} unique comments across {len(extractor.comments)} sections")
    print()

    # Generate and save output
    markdown_output = extractor.generate_markdown(school_name)
    output_path.write_text(markdown_output)

    print(f"Output saved to: {output_path}")


if __name__ == '__main__':
    main()
