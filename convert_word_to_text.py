from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


SUPPORTED_EXTENSIONS = {".docx", ".doc"}


def extract_docx_text(doc_path: Path) -> str:
    """Extract readable text from a DOCX file, including tables."""
    document = Document(doc_path)
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines) + "\n"


def convert_file(word_file: Path, output_folder: Path, finished_folder: Path) -> None:
    """Convert a single Word document to plain text and move the original."""
    suffix = word_file.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        print(f"Skipping unsupported file: {word_file.name}")
        return

    if suffix != ".docx":
        print(f"Cannot convert legacy Word file (DOC): {word_file.name}")
        return

    try:
        text_content = extract_docx_text(word_file)
        text_path = output_folder / word_file.with_suffix(".txt").name
        text_path.write_text(text_content, encoding="utf-8")
        destination = finished_folder / word_file.name
        shutil.move(str(word_file), destination)
        print(f"Converted {word_file.name} -> {text_path.name}")
    except Exception as exc:  # noqa: BLE001
        print(f"Failed to convert {word_file.name}: {exc}")


def main() -> None:
    print("Word to Plain Text Converter")
    print("============================")
    input_path = input("Enter path to folder containing Word documents: ").strip().strip('"')

    if not input_path:
        print("No path provided. Exiting.")
        return

    target_folder = Path(input_path).expanduser().resolve()
    if not target_folder.is_dir():
        print(f"Provided path is not a directory: {target_folder}")
        return

    word_files = [p for p in target_folder.iterdir() if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS]
    if not word_files:
        print("No Word documents (.docx or .doc) found in the provided folder.")
        return

    finished_folder = target_folder / "FinishedConverting"
    finished_folder.mkdir(exist_ok=True)

    for word_file in word_files:
        convert_file(word_file, target_folder, finished_folder)

    print("Done.")


if __name__ == "__main__":
    main()
