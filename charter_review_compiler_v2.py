#!/usr/bin/env python3
"""
Charter School Evaluation Comments Compiler - Version 2
Extracts reviewer comments from Word documents using XML parsing for form fields.
Combines XML table extraction with python-docx for section detection.
"""

import os
import sys
from pathlib import Path
from collections import defaultdict
from typing import Dict, List, Set, Tuple
from docx import Document
import zipfile
import xml.etree.ElementTree as ET
import tempfile


# Namespace mappings for Office Open XML
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
}

# Section lists for different application types
SECTION_LISTS = {
    '1': {  # Standard Application
        'name': 'Standard Application',
        'sections': [
            'Mission Guiding Principles and Purpose',
            'Target Population and Student Body',
            'Educational Program Design',
            'Curriculum and Instructional Design',
            'Student Performance',
            'Exceptional Students',
            'English Language Learners',
            'School Culture and Discipline',
            'Supplemental Programming',
            'Governance',
            'Management and Staffing',
            'Human Resources and Employment',
            'Professional Development',
            'Student Recruitment and Enrollment',
            'Parent and Community Involvement',
            'Facilities',
            'Transportation Service',
            'Food Service',
            'School Safety and Security',
            'Budget',
            'Financial Management and Oversight',
            'Start-Up Plan',
        ]
    },
    '2': {  # Virtual Application
        'name': 'Virtual Application',
        'sections': [
            'Mission, Guiding Principles and Purpose',
            'Target Population and Student Body',
            'Educational Program Design',
            'Curriculum Plan',
            'Student Performance, Assessment and Evaluation',
            'Exceptional Students',
            'English Language Learners',
            'School Culture and Discipline',
            'Supplemental Programming',
            'Governance',
            'Management and Staffing',
            'Human Resources and Employment',
            'Professional Development',
            'Student Recruitment and Enrollment',
            'Parent and Community Involvement',
            'Budget',
            'Financial Management and Oversight',
            'Start-Up Plan',
        ]
    },
    '3': {  # High Performing System Replication
        'name': 'High Performing System Replication',
        'sections': [
            'Replication Overview',
            'Mission Guiding Principles and Purpose',
            'Educational Program, Curriculum, and Instructional Design',
            'Student Performance',
            'Student Recruitment and Enrollment',
            'Management and Staffing',
            'Facilities',
            'Transportation Service',
            'Food Service',
            'School Safety and Security',
            'Budget',
            'Financial Management and Oversight',
        ]
    }
}


def extract_text_from_xml_element(element):
    """Extract all text from an XML element, including nested text nodes."""
    texts = []
    # Get direct text
    if element.text:
        texts.append(element.text)
    # Get text from all children
    for child in element.iter():
        if child.text and child != element:
            texts.append(child.text)
        if child.tail:
            texts.append(child.tail)
    return ''.join(texts).strip()


def extract_table_data_from_xml(doc_xml_path):
    """Extract table data from Word document XML, including content controls."""
    try:
        tree = ET.parse(doc_xml_path)
        root = tree.getroot()

        tables_data = []

        # Find all tables
        for table in root.findall('.//w:tbl', NAMESPACES):
            table_rows = []

            # Find all rows in the table
            for row in table.findall('.//w:tr', NAMESPACES):
                row_cells = []

                # Find all cells in the row
                for cell in row.findall('.//w:tc', NAMESPACES):
                    # Extract text from the cell, including content controls (sdt)
                    cell_text_parts = []

                    # Look for regular paragraphs
                    for para in cell.findall('.//w:p', NAMESPACES):
                        para_text = extract_text_from_xml_element(para)
                        if para_text:
                            cell_text_parts.append(para_text)

                    # Look for content controls (sdt elements) which may contain form fields
                    for sdt in cell.findall('.//w:sdt', NAMESPACES):
                        sdt_text = extract_text_from_xml_element(sdt)
                        if sdt_text and sdt_text not in cell_text_parts:
                            cell_text_parts.append(sdt_text)

                    cell_text = '\n'.join(cell_text_parts).strip()
                    row_cells.append(cell_text)

                if row_cells:  # Only add non-empty rows
                    table_rows.append(row_cells)

            if table_rows:
                tables_data.append(table_rows)

        return tables_data
    except Exception as e:
        print(f"  Warning: Error parsing XML: {e}")
        return []


def extract_tables_from_docx(docx_path):
    """Process a .docx file by extracting its XML and parsing tables."""
    try:
        # .docx files are zip archives
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            # Extract document.xml to temp location
            with tempfile.TemporaryDirectory() as temp_dir:
                # Extract the main document XML
                xml_content = zip_ref.read('word/document.xml')
                xml_path = Path(temp_dir) / 'document.xml'
                xml_path.write_bytes(xml_content)

                # Parse the XML
                return extract_table_data_from_xml(xml_path)
    except Exception as e:
        print(f"  Warning: Error processing DOCX: {e}")
        return []


class BoilerplateFilter:
    """Manages boilerplate text extraction from templates."""

    def __init__(self):
        self.boilerplate_texts: Set[str] = set()
        self.boilerplate_lines: Set[str] = set()

    def load_templates(self, template_folder: Path) -> int:
        """Load all template files and extract boilerplate text."""
        template_files = list(template_folder.glob('*.docx'))
        template_files = [f for f in template_files if not f.name.startswith('~$')]

        for template_file in template_files:
            try:
                doc = Document(template_file)

                # Extract all paragraph text
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        self.boilerplate_texts.add(text)
                        # Also store individual lines
                        for line in text.split('\n'):
                            clean_line = line.strip()
                            if clean_line:
                                self.boilerplate_lines.add(clean_line)

                # Extract all table text using XML to get form field data
                table_data = extract_tables_from_docx(template_file)
                for table in table_data:
                    for row in table:
                        for cell_text in row:
                            if cell_text:
                                self.boilerplate_texts.add(cell_text)
                                for line in cell_text.split('\n'):
                                    clean_line = line.strip()
                                    if clean_line:
                                        self.boilerplate_lines.add(clean_line)

            except Exception as e:
                print(f"  Warning: Could not load template {template_file.name}: {e}")

        return len(template_files)

    def is_boilerplate(self, text: str) -> bool:
        """Check if text is boilerplate."""
        text = text.strip()
        if not text:
            return True

        # Check exact match
        if text in self.boilerplate_texts or text in self.boilerplate_lines:
            return True

        # Check if it's a common header
        lower_text = text.lower()
        if lower_text in ['strengths', 'concerns and additional questions',
                          'concerns', 'reference', 'references',
                          'meets the standard', 'partially meets the standard',
                          'does not meet the standard', 'type', 'comment',
                          'strength', 'concern', 'question', 'follow up']:
            return True

        return False


class CommentExtractor:
    """Extracts comments from evaluation documents."""

    def __init__(self, boilerplate_filter: BoilerplateFilter, section_names: List[str]):
        self.boilerplate_filter = boilerplate_filter
        self.section_names = section_names
        self.comments: Dict[str, Dict[str, List[Tuple[str, str]]]] = defaultdict(
            lambda: {'strengths': [], 'concerns': []}
        )
        self.boilerplate_filtered_count = 0

    def extract_reviewer_name(self, filename: str) -> str:
        """Extract reviewer name from filename."""
        parts = filename.replace('.docx', '').split('_')
        if len(parts) >= 3:
            # Handle both "Eval" and "Evanl" (typo in filenames)
            for i, part in enumerate(parts):
                if part.lower() in ['eval', 'evanl']:
                    if i + 1 < len(parts):
                        return parts[i + 1]
        return 'Unknown Reviewer'

    def extract_school_name(self, filename: str) -> str:
        """Extract school name from filename."""
        parts = filename.replace('.docx', '').split('_')
        if len(parts) >= 3:
            for i, part in enumerate(parts):
                if part.lower() in ['eval', 'evanl']:
                    return '_'.join(parts[:i])
        return 'Unknown School'

    def find_section_name(self, text: str) -> str:
        """Find matching section name from text."""
        text = text.strip()

        # Try to match against known section names
        for section_name in self.section_names:
            # Check if section name appears in text (case-insensitive)
            if section_name.lower() in text.lower():
                return section_name

            # Check if text appears in section name
            if text.lower() in section_name.lower():
                return section_name

        # Check for addendum sections
        if 'addendum' in text.lower():
            return text

        return None

    def parse_table_for_comments(self, table_data: List[List[str]]) -> Tuple[List[str], List[str]]:
        """Parse a table (as list of rows) to extract strengths and concerns.

        Handles two formats:
        1. Type | Reference | Comment (FCI-branded format)
        2. Multi-row format with Strengths/Concerns headers (State model)
        """
        strengths = []
        concerns = []

        if not table_data or len(table_data) == 0:
            return strengths, concerns

        # Check if this is a Type|Reference|Comment format table
        if len(table_data) > 0 and len(table_data[0]) >= 3:
            first_row_cells = [cell.lower() for cell in table_data[0]]

            # Check if this is a Type|Reference|Comment header row
            if 'type' in first_row_cells[0] and ('reference' in first_row_cells[1] or 'comment' in ' '.join(first_row_cells)):
                # This is the Type|Reference|Comment format
                # Process rows after the header
                for row_idx in range(1, len(table_data)):
                    row = table_data[row_idx]
                    if len(row) < 2:
                        continue

                    # Skip empty rows
                    if not any(row):
                        continue

                    # Get type, reference, and comment
                    comment_type = row[0].lower() if len(row) > 0 else ''
                    reference = row[1] if len(row) > 1 else ''
                    comment = row[2] if len(row) > 2 else row[1] if len(row) > 1 else ''

                    # Skip if comment is empty or boilerplate
                    if not comment or self.boilerplate_filter.is_boilerplate(comment):
                        if comment:
                            self.boilerplate_filtered_count += 1
                        continue

                    # Add page reference if available
                    comment_text = comment
                    if reference and not self.boilerplate_filter.is_boilerplate(reference):
                        # Check if reference looks like a page number
                        if reference.strip().replace('.', '').replace('p', '').replace(' ', '').isdigit():
                            # Format as [p. XX]
                            page_num = reference.strip().replace('p.', '').replace('p', '').strip()
                            comment_text = f"{comment} [p. {page_num}]"
                        elif reference and len(reference) < 10:  # Short reference, likely a page number
                            comment_text = f"{comment} [{reference}]"

                    # Categorize by type
                    if 'strength' in comment_type:
                        strengths.append(comment_text)
                    elif 'concern' in comment_type or 'question' in comment_type:
                        concerns.append(comment_text)
                    # Note: "Follow Up" type goes to concerns by default

                return strengths, concerns

        # Fall back to original multi-row format parsing (State model)
        current_section = None

        for row_idx, row in enumerate(table_data):
            first_cell = row[0].lower() if row and len(row) > 0 else ''

            if 'strength' in first_cell:
                current_section = 'strengths'
                continue
            elif 'concern' in first_cell or 'question' in first_cell:
                current_section = 'concerns'
                continue

            # Extract content from cells
            for cell_idx, cell_text in enumerate(row):
                if not cell_text or self.boilerplate_filter.is_boilerplate(cell_text):
                    if cell_text:
                        self.boilerplate_filtered_count += 1
                    continue

                if cell_text.lower() in ['reference', 'references', 'strengths',
                                   'concerns', 'concerns and additional questions']:
                    continue

                comment_text = cell_text
                # Skip if it looks like a page reference in the last column
                if cell_idx == len(row) - 1 and cell_idx > 0:
                    if cell_text.startswith('p.') or cell_text.isdigit():
                        continue

                if current_section == 'strengths':
                    strengths.append(comment_text)
                elif current_section == 'concerns':
                    concerns.append(comment_text)

        return strengths, concerns

    def process_document(self, file_path: Path) -> int:
        """Process a single evaluation document."""
        try:
            # Use python-docx for paragraph reading (section detection)
            doc = Document(file_path)
            reviewer_name = self.extract_reviewer_name(file_path.name)
            comment_count = 0

            # Build section mapping from paragraphs
            section_sequence = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Check if this is a section header
                section_name = self.find_section_name(text)
                if section_name:
                    section_sequence.append(section_name)

            # Use XML parsing for table extraction
            table_data_list = extract_tables_from_docx(file_path)

            # Match tables to sections based on sequence
            current_section_idx = 0

            for table_data in table_data_list:
                strengths, concerns = self.parse_table_for_comments(table_data)

                # Assign to current section if we have one
                if current_section_idx < len(section_sequence):
                    current_section = section_sequence[current_section_idx]
                    # Advance to next section if we found comments
                    if strengths or concerns:
                        current_section_idx += 1
                else:
                    # If we have comments but no section, use generic
                    if strengths or concerns:
                        current_section = "General Comments"
                    else:
                        continue

                # Add comments to the section
                for strength in strengths:
                    if not self.boilerplate_filter.is_boilerplate(strength):
                        self.comments[current_section]['strengths'].append(
                            (reviewer_name, strength)
                        )
                        comment_count += 1
                    else:
                        self.boilerplate_filtered_count += 1

                for concern in concerns:
                    if not self.boilerplate_filter.is_boilerplate(concern):
                        self.comments[current_section]['concerns'].append(
                            (reviewer_name, concern)
                        )
                        comment_count += 1
                    else:
                        self.boilerplate_filtered_count += 1

            return comment_count

        except Exception as e:
            print(f"  ✗ Error processing {file_path.name}: {e}")
            import traceback
            traceback.print_exc()
            return 0

    def remove_duplicates(self) -> int:
        """Remove exact duplicate comments."""
        duplicates_removed = 0

        for section in self.comments:
            # Remove duplicates from strengths
            seen = set()
            unique_strengths = []
            for reviewer, comment in self.comments[section]['strengths']:
                if comment not in seen:
                    seen.add(comment)
                    unique_strengths.append((reviewer, comment))
                else:
                    duplicates_removed += 1
            self.comments[section]['strengths'] = unique_strengths

            # Remove duplicates from concerns
            seen = set()
            unique_concerns = []
            for reviewer, comment in self.comments[section]['concerns']:
                if comment not in seen:
                    seen.add(comment)
                    unique_concerns.append((reviewer, comment))
                else:
                    duplicates_removed += 1
            self.comments[section]['concerns'] = unique_concerns

        return duplicates_removed

    def generate_markdown(self, school_name: str) -> str:
        """Generate Markdown output."""
        lines = []
        lines.append("# Charter Application Review Comments Compilation")
        lines.append(f"## School Name: {school_name}")
        lines.append("")

        # Process sections in order
        for idx, section_name in enumerate(self.section_names, 1):
            if section_name not in self.comments:
                continue

            section_data = self.comments[section_name]

            # Skip if no comments in this section
            if not section_data['strengths'] and not section_data['concerns']:
                continue

            lines.append(f"## Section {idx}. {section_name}")

            # Add strengths
            if section_data['strengths']:
                lines.append("### Strengths")
                for reviewer, comment in section_data['strengths']:
                    lines.append(f"- {reviewer}: {comment}")
                lines.append("")

            # Add concerns
            if section_data['concerns']:
                lines.append("### Concerns")
                for reviewer, comment in section_data['concerns']:
                    lines.append(f"- {reviewer}: {comment}")
                lines.append("")

        # Add any addendum sections
        for section_name in self.comments:
            if section_name not in self.section_names and section_name != "General Comments":
                section_data = self.comments[section_name]

                if not section_data['strengths'] and not section_data['concerns']:
                    continue

                lines.append(f"## {section_name}")

                if section_data['strengths']:
                    lines.append("### Strengths")
                    for reviewer, comment in section_data['strengths']:
                        lines.append(f"- {reviewer}: {comment}")
                    lines.append("")

                if section_data['concerns']:
                    lines.append("### Concerns")
                    for reviewer, comment in section_data['concerns']:
                        lines.append(f"- {reviewer}: {comment}")
                    lines.append("")

        return '\n'.join(lines)


def main():
    """Main program execution."""
    print("Charter School Evaluation Comments Compiler - Version 2")
    print("=" * 55)
    print("Using XML parsing for form field extraction")
    print()

    # Prompt for review documents folder
    review_folder_path = input("Enter path to folder containing review documents: ").strip()
    review_folder = Path(review_folder_path)

    if not review_folder.exists() or not review_folder.is_dir():
        print(f"Error: Folder not found: {review_folder_path}")
        sys.exit(1)

    docx_files = [f for f in review_folder.glob('*.docx') if not f.name.startswith('~$')]

    if not docx_files:
        print(f"Error: No .docx files found in {review_folder_path}")
        sys.exit(1)

    print(f"Found {len(docx_files)} .docx file(s)")
    print()

    # Prompt for template folder
    template_folder_path = input("Enter path to folder containing evaluation matrix templates: ").strip()
    template_folder = Path(template_folder_path)

    if not template_folder.exists() or not template_folder.is_dir():
        print(f"Error: Folder not found: {template_folder_path}")
        sys.exit(1)

    # Load templates
    print("Loading templates...")
    boilerplate_filter = BoilerplateFilter()
    template_count = boilerplate_filter.load_templates(template_folder)
    print(f"Loaded {template_count} template file(s)")
    print()

    # Prompt for application type
    print("Select application type:")
    print("  1. Standard Application (22 sections)")
    print("  2. Virtual Application (18 sections)")
    print("  3. High Performing System Replication (12 sections)")

    app_type = input("Enter selection (1-3): ").strip()

    if app_type not in SECTION_LISTS:
        print(f"Error: Invalid selection: {app_type}")
        sys.exit(1)

    section_info = SECTION_LISTS[app_type]
    print()

    # Extract school name from first file
    extractor_temp = CommentExtractor(boilerplate_filter, section_info['sections'])
    school_name = extractor_temp.extract_school_name(docx_files[0].name)

    # Prompt for output filename
    default_output = f"{school_name}_CompiledReviews.md"
    output_filename = input(f"Enter output filename [{default_output}]: ").strip()
    if not output_filename:
        output_filename = default_output

    output_path = review_folder / output_filename
    print()

    # Process documents
    print("Processing reviews...")
    extractor = CommentExtractor(boilerplate_filter, section_info['sections'])

    for docx_file in docx_files:
        comment_count = extractor.process_document(docx_file)
        print(f"  ✓ {docx_file.name} ({comment_count} comments extracted)")

    print()

    # Remove duplicates
    duplicates = extractor.remove_duplicates()

    # Calculate total unique comments
    total_comments = sum(
        len(section['strengths']) + len(section['concerns'])
        for section in extractor.comments.values()
    )

    print(f"Removed {extractor.boilerplate_filtered_count} boilerplate text instances")
    print(f"Removed {duplicates} duplicate comments")
    print(f"Compiled {total_comments} unique comments across {len(extractor.comments)} sections")
    print()

    # Generate and save output
    markdown_output = extractor.generate_markdown(school_name)
    output_path.write_text(markdown_output, encoding='utf-8')

    print(f"Output saved to: {output_path}")


if __name__ == '__main__':
    main()
