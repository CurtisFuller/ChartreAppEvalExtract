#!/usr/bin/env python3
"""
Debug script to show exactly what text is being extracted from Word files.
This will help us see if the form field data is accessible.
"""

import sys
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("ERROR: python-docx not installed. Install with: pip install python-docx")
    sys.exit(1)


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract all text from a Word document."""
    try:
        doc = Document(docx_path)
        text_parts = []

        # Extract all paragraphs and tables in document order
        for element in doc.element.body:
            # Check if it's a paragraph
            if element.tag.endswith('p'):
                for para in doc.paragraphs:
                    if para._element == element:
                        para_text = para.text.strip()
                        if para_text:
                            text_parts.append(para_text)
                        break

            # Check if it's a table
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        # Extract table content with tabs between cells
                        for row in table.rows:
                            row_cells = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                row_cells.append(cell_text)

                            # Join cells with tabs
                            if any(row_cells):
                                text_parts.append('\t'.join(row_cells))
                        break

        return '\n'.join(text_parts)

    except Exception as e:
        raise Exception(f"Error extracting text from {docx_path.name}: {e}")


def main():
    """Main execution."""
    print("Word Document Text Extraction Debug Tool")
    print("=" * 45)
    print()

    # Get file path
    file_path = input("Enter path to .docx file: ").strip()
    docx_file = Path(file_path)

    if not docx_file.exists():
        print(f"ERROR: File not found: {file_path}")
        sys.exit(1)

    if docx_file.suffix.lower() != '.docx':
        print(f"ERROR: File must be a .docx file")
        sys.exit(1)

    print(f"\nExtracting text from: {docx_file.name}")
    print("-" * 70)
    print()

    # Extract text
    try:
        extracted_text = extract_text_from_docx(docx_file)

        # Save to file
        output_file = docx_file.with_suffix('.extracted.txt')
        output_file.write_text(extracted_text, encoding='utf-8')

        print(f"✓ Text extracted successfully!")
        print(f"✓ Saved to: {output_file}")
        print()
        print("=" * 70)
        print("EXTRACTED TEXT PREVIEW (first 3000 characters):")
        print("=" * 70)
        print(extracted_text[:3000])
        print()
        print("=" * 70)
        print(f"Total characters extracted: {len(extracted_text)}")
        print(f"Total lines: {len(extracted_text.split(chr(10)))}")
        print()

        # Look for comment sections
        import re
        comment_sections = re.findall(r'Section\s+\d+.*?Evaluation\s+Comments:', extracted_text, re.IGNORECASE)
        print(f"Found {len(comment_sections)} section evaluation comment blocks")
        if comment_sections:
            print("\nSection headers found:")
            for section in comment_sections[:5]:  # Show first 5
                print(f"  - {section}")

        # Look for table data patterns
        table_rows = [line for line in extracted_text.split('\n') if '\t' in line]
        print(f"\nFound {len(table_rows)} lines with tab-separated data (table rows)")

        # Show some table rows
        if table_rows:
            print("\nSample table rows (first 10):")
            for i, row in enumerate(table_rows[:10], 1):
                parts = row.split('\t')
                print(f"  Row {i}: {len(parts)} columns - {parts[:3]}")

    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
